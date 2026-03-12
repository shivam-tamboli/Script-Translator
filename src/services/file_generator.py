from pathlib import Path
from typing import Optional
from docx import Document
from docx.shared import Pt


class FileGenerationError(Exception):
    """Raised when file generation fails."""
    pass


class FileGenerator:
    def generate_docx(self, text: str, output_path: str) -> str:
        """Generate a DOCX file with translated text."""
        try:
            doc = Document()
            
            heading = doc.add_heading("Translated Text", level=1)
            heading.runs[0].font.size = Pt(16)
            
            for paragraph in text.split("\n"):
                if paragraph.strip():
                    doc.add_paragraph(paragraph)
            
            doc.save(output_path)
            return output_path
        except Exception as e:
            raise FileGenerationError(f"Failed to generate DOCX: {str(e)}")

    def generate(
        self,
        text: str,
        output_path: str,
        format: Optional[str] = None
    ) -> str:
        """Generate output file based on extension or format."""
        path = Path(output_path)
        extension = format or path.suffix.lower()
        
        if extension in [".docx", ".doc"]:
            return self.generate_docx(text, output_path)
        else:
            raise FileGenerationError(f"Unsupported output format: {extension}")
